"""Build a DOCX guide for approaching the SportRadar Tennis API project."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path("SportRadar_Tennis_API_Smart_Guide.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 32, 46)
MUTED = RGBColor(91, 105, 120)
LIGHT_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_pdf_summary(doc)
    add_endpoint_strategy(doc)
    add_request_pattern(doc)
    add_etl_workflow(doc)
    add_database_and_queries(doc)
    add_streamlit_plan(doc)
    add_image_analysis_plan(doc)
    add_testing_and_limits(doc)
    add_execution_checklist(doc)
    add_sources(doc)
    doc.save(OUTPUT)
    print(f"created {OUTPUT}")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("SportRadar Tennis API Smart Guide")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("SportRadar Tennis API Smart Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Project approach for Game Analytics: Unlocking Tennis Data with SportRadar API"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Purpose",
        "This document translates the supplied 14-page project PDF into an implementation-ready plan: how to request the SportRadar Tennis API, flatten the responses into SQL tables, build analytics in Streamlit, and add an image-based analysis extension using local/free tooling.",
    )


def add_pdf_summary(doc: Document) -> None:
    doc.add_heading("1. What the PDF Requires", level=1)
    add_bullets(
        doc,
        [
            "Build a Python, SQL, and Streamlit analytics project around SportRadar Tennis data.",
            "Collect three feeds: Competitions, Complexes, and Doubles Competitor Rankings.",
            "Normalize the data into six tables: Categories, Competitions, Complexes, Venues, Competitors, and Competitor_Rankings.",
            "Write analytical SQL queries for competitions, venues, and competitors.",
            "Create an interactive Streamlit dashboard with filters, tables, charts, competitor details, country analysis, and leaderboards.",
            "Document setup, workflow, SQL queries, challenges, and insights for GitHub submission.",
        ],
    )

    add_table(
        doc,
        ["PDF Area", "Implementation Meaning"],
        [
            ["Data extraction", "Use authenticated GET requests to official SportRadar Tennis v3 endpoints."],
            ["Data storage", "Create normalized SQL tables with primary keys and foreign keys."],
            ["Data analysis", "Expose SQL query results through Pandas DataFrames."],
            ["Streamlit app", "Use filters, pagination, charts, and detail views over the database."],
            ["Evaluation", "Show API accuracy, database design, query efficiency, usability, error handling, and documentation quality."],
        ],
        [2.0, 4.5],
    )


def add_endpoint_strategy(doc: Document) -> None:
    doc.add_heading("2. SportRadar API Request Strategy", level=1)
    doc.add_paragraph(
        "SportRadar Tennis v3 uses URL path parameters for access level, language, and response format. The three project feeds are all GET endpoints. Use JSON format and keep credentials in environment variables."
    )
    add_table(
        doc,
        ["Feed", "Official endpoint pattern", "Purpose"],
        [
            [
                "Competitions",
                "GET /tennis/{access_level}/v3/{language_code}/competitions.{format}",
                "Returns all competitions; transform categories and competitions.",
            ],
            [
                "Complexes",
                "GET /tennis/{access_level}/v3/{language_code}/complexes.{format}",
                "Returns complexes and nested venues; transform complexes and venues.",
            ],
            [
                "Doubles rankings",
                "GET /tennis/{access_level}/v3/{language_code}/double_competitors_rankings.{format}",
                "Returns top 500 ATP/WTA doubles rankings; transform competitors and ranking rows.",
            ],
        ],
        [1.45, 3.15, 1.9],
    )

    doc.add_heading("Recommended environment variables", level=2)
    add_table(
        doc,
        ["Variable", "Example", "Why it matters"],
        [
            ["SPORTRADAR_API_KEY", "your-key", "Required credential; never hard-code it."],
            ["SPORTRADAR_ACCESS_LEVEL", "trial", "Use trial while developing; production only if your account supports it."],
            ["SPORTRADAR_LANGUAGE_CODE", "en", "Controls localized response text."],
            ["DATABASE_URL", "postgresql+psycopg://...", "Lets the same ETL code target PostgreSQL or MySQL."],
        ],
        [1.8, 2.1, 2.6],
    )


def add_request_pattern(doc: Document) -> None:
    doc.add_heading("3. Python Request Pattern", level=1)
    add_callout(
        doc,
        "Rule",
        "Build one API client that owns base URL construction, headers, timeouts, retries, rate-limit handling, and JSON parsing. Do not scatter raw requests across the project.",
    )
    code = """import os
import requests

BASE_URL = "https://api.sportradar.com/tennis"
API_KEY = os.getenv("SPORTRADAR_API_KEY")
ACCESS_LEVEL = os.getenv("SPORTRADAR_ACCESS_LEVEL", "trial")
LANGUAGE = os.getenv("SPORTRADAR_LANGUAGE_CODE", "en")

def get_feed(endpoint: str) -> dict:
    if not API_KEY:
        raise ValueError("SPORTRADAR_API_KEY is missing")
    url = f"{BASE_URL}/{ACCESS_LEVEL}/v3/{LANGUAGE}/{endpoint}.json"
    response = requests.get(
        url,
        headers={"x-api-key": API_KEY},
        timeout=20,
    )
    if response.status_code == 429:
        raise RuntimeError("Rate limit reached; wait before retrying")
    response.raise_for_status()
    return response.json()

competitions = get_feed("competitions")
complexes = get_feed("complexes")
rankings = get_feed("double_competitors_rankings")"""
    add_code_block(doc, code)

    doc.add_heading("Response handling checklist", level=2)
    add_bullets(
        doc,
        [
            "Validate that the response root is a JSON object before transforming it.",
            "For 401 errors, check the API key, account access level, and header name.",
            "For 429 errors, honor Retry-After when present and avoid rapid repeated calls.",
            "For 5xx errors, retry with exponential backoff and log the endpoint name.",
            "Keep raw response samples during development, but do not commit real API credentials.",
        ],
    )


def add_etl_workflow(doc: Document) -> None:
    doc.add_heading("4. Transform and Load Workflow", level=1)
    add_numbered(
        doc,
        [
            "Fetch competitions, complexes, and double_competitors_rankings from the API.",
            "Flatten nested JSON objects into table-shaped rows.",
            "Upsert categories before competitions, complexes before venues, and competitors before rankings.",
            "Run SQL analytical queries through wrapper functions that return Pandas DataFrames.",
            "Render the Streamlit dashboard from query outputs, not from raw JSON.",
        ],
    )
    add_table(
        doc,
        ["Source feed", "Nested source objects", "Target tables"],
        [
            ["competitions", "competition.category, competition.parent", "Categories, Competitions"],
            ["complexes", "complex.venues[]", "Complexes, Venues"],
            ["double_competitors_rankings", "rankings[].competitor_rankings[].competitor", "Competitors, Competitor_Rankings"],
        ],
        [1.55, 2.4, 2.55],
    )


def add_database_and_queries(doc: Document) -> None:
    doc.add_heading("5. Database and Query Layer", level=1)
    doc.add_paragraph(
        "Use SportRadar string IDs, such as sr:competition:123, as stable primary keys for entity tables. Use an auto-increment rank_id for competitor rankings and keep competitor_id as the foreign key."
    )
    add_table(
        doc,
        ["Query group", "Examples required by the PDF"],
        [
            ["Competitions", "Categories, doubles filter, category filter, parent/sub-competition relationships, type distribution, top-level competitions."],
            ["Venues", "Complex joins, venues per complex, country filter, timezones, multi-venue complexes, country grouping, venues for a named complex."],
            ["Competitors", "Rank and points, top 5, stable movement, country points, country counts, highest current points."],
        ],
        [1.55, 4.95],
    )
    add_callout(
        doc,
        "Implementation status",
        "The project workspace already contains a Phase 1 ETL package, Phase 2 query wrappers, and Phase 3-4 Streamlit views. The missing production step is running the ETL with a real SportRadar API key and a live PostgreSQL/MySQL database.",
    )


def add_streamlit_plan(doc: Document) -> None:
    doc.add_heading("6. Streamlit Dashboard Plan", level=1)
    add_bullets(
        doc,
        [
            "Home Dashboard: metric cards and chart grid for points, ranks, and countries.",
            "Filtered Rankings: searchable table with rank range and country filters plus pagination.",
            "Competitor Details: selected competitor profile with rank, points, movement, and competitions played.",
            "Country-Wise Analysis: country table and choropleth map using ISO country codes.",
            "Leaderboards: podium blocks for top rank, most points, and most improved.",
        ],
    )
    add_callout(
        doc,
        "Data limitation",
        "The PDF schema does not include age, turned-pro date, career titles, win/loss records, or historical weekly ranking snapshots. If those are required, add new SportRadar feeds and extend the database schema before showing those fields as real values.",
    )


def add_image_analysis_plan(doc: Document) -> None:
    doc.add_heading("7. Adding Image-Based Analysis", level=1)
    doc.add_paragraph(
        "The PDF focuses on API, SQL, and Streamlit, but an image-based analysis layer can make the project more impressive. Keep it optional and local-first so it stays within free limits."
    )

    add_table(
        doc,
        ["Image feature", "Free/local implementation", "Dashboard output"],
        [
            [
                "Upload court or player image",
                "Streamlit file_uploader plus Pillow/OpenCV for metadata, dimensions, color, and quality checks.",
                "Image preview, size, aspect ratio, brightness, blur score.",
            ],
            [
                "Extract text from screenshots",
                "Tesseract OCR or a local OCR package if installed; otherwise provide manual text capture fallback.",
                "Detected player names/ranks from screenshots for comparison with database rows.",
            ],
            [
                "Chart screenshot QA",
                "Save dashboard screenshots and compare image dimensions/blankness locally.",
                "Automated check that dashboard charts are non-empty and readable.",
            ],
            [
                "Player/country image enrichment",
                "Store image paths or URLs in an optional Images table; do not store binary blobs unless necessary.",
                "Competitor profile image, source URL, alt text, and image status.",
            ],
        ],
        [1.6, 3.0, 1.9],
    )

    doc.add_heading("Optional schema extension", level=2)
    add_code_block(
        doc,
        """CREATE TABLE images (
    image_id SERIAL PRIMARY KEY,
    entity_type VARCHAR(40) NOT NULL,
    entity_id VARCHAR(64) NOT NULL,
    image_path TEXT NOT NULL,
    source_url TEXT,
    alt_text TEXT,
    width_px INT,
    height_px INT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);""",
    )

    doc.add_heading("Streamlit image analysis workflow", level=2)
    add_numbered(
        doc,
        [
            "Add a new sidebar link named Image Analysis.",
            "Use st.file_uploader to accept PNG/JPG images.",
            "Use Pillow to open the image and calculate size, aspect ratio, brightness, and sharpness.",
            "Show the uploaded image, computed metrics, and a warning if the image is too small or blurry.",
            "Optionally map OCR-detected names to competitor rows from the database.",
        ],
    )


def add_testing_and_limits(doc: Document) -> None:
    doc.add_heading("8. Testing, Rate Limits, and Free-Tier Safety", level=1)
    add_bullets(
        doc,
        [
            "Run transform tests with saved sample JSON before calling live endpoints repeatedly.",
            "Cache API responses locally during development to avoid unnecessary requests.",
            "Call each SportRadar endpoint only when needed, then develop against the database.",
            "Handle 401, 429, and 5xx errors clearly in logs.",
            "Keep image analysis local with Pillow/OpenCV/OCR; avoid paid external image APIs unless explicitly required.",
        ],
    )
    add_table(
        doc,
        ["Test layer", "What to verify"],
        [
            ["API client", "URL construction, headers, timeout, 401/429/5xx handling."],
            ["Transforms", "Nested JSON becomes valid table rows with required keys."],
            ["Database", "Foreign keys, upserts, idempotent reruns."],
            ["Queries", "All required analytical queries return DataFrames."],
            ["Streamlit", "Filters, pagination, empty states, and chart rendering."],
            ["Images", "Upload validation, nonblank image checks, OCR fallback behavior."],
        ],
        [1.45, 5.05],
    )


def add_execution_checklist(doc: Document) -> None:
    doc.add_heading("9. Execution Checklist", level=1)
    add_numbered(
        doc,
        [
            "Create a SportRadar account and obtain the Tennis API key.",
            "Set SPORTRADAR_API_KEY, SPORTRADAR_ACCESS_LEVEL, SPORTRADAR_LANGUAGE_CODE, and DATABASE_URL.",
            "Run the schema creation script.",
            "Run the ETL once and confirm row counts in all six tables.",
            "Run the analytical query functions and save screenshots of successful results.",
            "Start Streamlit and verify all dashboard views.",
            "Add the image analysis view if you want an innovation/creativity enhancement.",
            "Upload code, SQL scripts, README, and documentation to GitHub.",
        ],
    )


def add_sources(doc: Document) -> None:
    doc.add_heading("10. Sources and Links", level=1)
    add_bullets(
        doc,
        [
            "Project PDF: Game Analytics: Unlocking Tennis Data with SportRadar API.",
            "SportRadar Competitions reference: https://developer.sportradar.com/tennis/reference/competitions",
            "SportRadar Complexes reference: https://developer.sportradar.com/tennis/reference/complexes",
            "SportRadar Doubles Competitor Rankings reference: https://developer.sportradar.com/tennis/reference/doubles-competitor-rankings",
            "SportRadar signup: https://console.sportradar.com/signup",
            "Streamlit API reference: https://docs.streamlit.io/library/api-reference",
        ],
    )


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = DARK_BLUE
    paragraph.add_run(text)
    doc.add_paragraph()


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade_cell(cell, "F2F4F7")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code.splitlines()):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 42, 56)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, widths)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        shade_cell(cell, LIGHT_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = DARK_BLUE

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.add_run(value)
    doc.add_paragraph()


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


if __name__ == "__main__":
    main()
